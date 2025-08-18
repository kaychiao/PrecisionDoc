"""
Export utilities for Word documents.
"""
import os
import pandas as pd
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from precisiondoc.utils.word.document_formatting import DocumentFormatter
from precisiondoc.utils.word.evidence_processing import EvidenceProcessor
from precisiondoc.utils.word.content_formatting import ContentFormatter
from precisiondoc.utils.word.table_utils import TableUtils
from precisiondoc.utils.word.image_utils import ImageUtils

class ExportUtils:
    """Handles document export operations"""
    
    @staticmethod
    def _create_evidence_table(doc, evidence_dict, image_path, multi_line_text=True, show_borders=True, exclude_columns=None):
        """
        Create a table for evidence with text and image placeholders
        
        Args:
            doc: Word document
            evidence_dict: Dictionary containing evidence data
            image_path: Path to image file
            multi_line_text: Whether to use multi-line text format (one row per key-value pair)
            show_borders: Whether to show table borders
            exclude_columns: Columns to exclude from evidence text
            
        Returns:
            table: Created table
        """
        if exclude_columns is None:
            exclude_columns = ['is_precision_evidence', 'page_number', 'document_type']
        
        # Rename text to resource_sentence, image_path to resource_url (key name replacement)
        evidence_dict = EvidenceProcessor._rename_evidence_dict_keys(evidence_dict)
        
        # Sort the evidence dictionary based on the predefined key sequence
        evidence_dict = EvidenceProcessor._sort_evidence_dict(evidence_dict)
        
        # Create table based on format (multi-line or single-row JSON)
        if multi_line_text:
            # Count non-excluded items with non-empty values
            row_count = sum(1 for k, v in evidence_dict.items() 
                           if k not in exclude_columns and v and str(v).strip())
        else:
            row_count = 1
        
        # Create basic table
        table = TableUtils._create_basic_table(doc, row_count)
        
        # Get current section and check if it's landscape
        current_section = doc.sections[-1]
        is_landscape = current_section.orientation == 1  # 1 is landscape
        
        # Set table column widths
        table_width, text_col_width, image_col_width = TableUtils._set_table_column_widths(table, is_landscape)
        
        # Set table pagination properties
        TableUtils._set_table_pagination_properties(table)
        
        # Populate table cells
        if multi_line_text:
            # Populate with one key-value pair per row
            left_cells, right_cell = EvidenceProcessor._populate_multi_line_cells(table, evidence_dict, exclude_columns)
        else:
            # Populate with JSON-style dictionary in a single cell
            left_cell, right_cell = EvidenceProcessor._populate_json_format_cell(table, evidence_dict)
        
        # Add image to the right column
        ImageUtils._add_image_to_cell(right_cell, image_path, image_col_width)
        
        # Remove table borders if specified
        if not show_borders:
            TableUtils._remove_table_borders(table)
            
        return table
    
    @staticmethod
    def _add_separator_and_section(doc):
        """
        Add a separator line and section break
        
        Args:
            doc: Word document
        """
        # Add separator line
        separator = doc.add_paragraph('─' * 100)
        ContentFormatter.apply_separator_format(separator)
        
        # Add section break for next evidence
        doc.add_section()
    
    @staticmethod
    def export_evidence_to_word(evidence_data, output_file, output_folder=None, multi_line_text=True, 
                               show_borders=True, exclude_columns=None, page_settings=None):
        """
        Export evidence data to a Word document.
        
        Args:
            evidence_data: Path to Excel file or pandas DataFrame with evidence data
            output_file: Path to output Word file
            output_folder: Folder containing images referenced in evidence data
            multi_line_text: Whether to use multi-line text format (True) or JSON format (False)
            show_borders: Whether to show borders in tables
            exclude_columns: List of columns to exclude from evidence text
            page_settings: Dictionary with page settings for Word document.
                          Supported keys: 'orientation' ('portrait' or 'landscape'),
                          'margins' (dict with 'left', 'right', 'top', 'bottom' in inches)
        """
        logger = logging.getLogger(__name__)
        
        try:
            # Handle both DataFrame and file path inputs
            if isinstance(evidence_data, str):
                if evidence_data.endswith('.xlsx') or evidence_data.endswith('.xls'):
                    df = pd.read_excel(evidence_data)
                else:
                    raise ValueError(f"Unsupported file format: {evidence_data}")
            elif isinstance(evidence_data, pd.DataFrame):
                df = evidence_data
            else:
                raise ValueError("evidence_data must be a DataFrame or path to an Excel file")
            
            # Filter for precision evidence only if the column exists
            if 'is_precision_evidence' in df.columns:
                # Handle various true value formats (True, 'true', 'True', 1, etc.)
                true_values = [True, 'true', 'True', 'TRUE', 1, '1', 'yes', 'Yes', 'YES']
                df = df[df['is_precision_evidence'].astype(str).isin([str(v) for v in true_values])]
            
            if df.empty:
                logger.warning("No evidence data to export to Word")
                return
                
            # Create Word document
            doc = Document()
            
            # Apply document formatting
            DocumentFormatter.apply_word_document_format(doc)
            
            # Get default orientation from page_settings or use landscape as default
            default_orientation = 'landscape'
            if page_settings and 'orientation' in page_settings:
                default_orientation = page_settings['orientation'].lower()
            
            # Ensure the first section has the correct orientation
            first_section = doc.sections[0]
            DocumentFormatter.set_section_orientation(first_section, default_orientation)
            
            # Apply custom margins if specified
            if page_settings and 'margins' in page_settings:
                margins = page_settings['margins']
                if 'left' in margins:
                    first_section.left_margin = Inches(margins['left'])
                if 'right' in margins:
                    first_section.right_margin = Inches(margins['right'])
                if 'top' in margins:
                    first_section.top_margin = Inches(margins['top'])
                if 'bottom' in margins:
                    first_section.bottom_margin = Inches(margins['bottom'])
            
            # Define default columns to exclude from evidence text if not provided
            if exclude_columns is None:
                exclude_columns = (
                    'page_type', 'is_precision_evidence',
                    'page_number', 'success', 'document_type', '解析', '分析', '结论', '文字提取', 
                    'evidence_level', 'evidence_type', 'evidence_list'
                )
            
            # Process evidence row by row
            for idx, row in df.iterrows():
                # Create a new section for each evidence item (except the first one)
                if idx > 0:
                    doc.add_section()
                    section = doc.sections[idx]
                    # Apply the same orientation to all sections
                    DocumentFormatter.set_section_orientation(section, default_orientation)
                    
                    # Apply custom margins if specified
                    if page_settings and 'margins' in page_settings:
                        margins = page_settings['margins']
                        if 'left' in margins:
                            section.left_margin = Inches(margins['left'])
                        if 'right' in margins:
                            section.right_margin = Inches(margins['right'])
                        if 'top' in margins:
                            section.top_margin = Inches(margins['top'])
                        if 'bottom' in margins:
                            section.bottom_margin = Inches(margins['bottom'])
                
                # Prepare evidence text
                evidence_dict = EvidenceProcessor._prepare_evidence_text(row, exclude_columns)
                
                # Get image path
                image_path = row.get('image_path', '')
                
                # Create table with text and image placeholders
                ExportUtils._create_evidence_table(doc, evidence_dict, image_path, 
                                                  multi_line_text=multi_line_text,
                                                  show_borders=show_borders,
                                                  exclude_columns=exclude_columns)
                
                # Add separator line
                separator = doc.add_paragraph('-' * 50)
                ContentFormatter.apply_separator_format(separator)
            
            # Create directory if it doesn't exist
            if output_file is not None:
                os.makedirs(os.path.dirname(output_file), exist_ok=True)
                
                # Save Word document
                doc.save(output_file)
                logger.info(f"Evidence exported to Word file: {output_file}")
            else:
                logger.error("No output file path provided and could not generate a default path")
            return output_file
        except Exception as e:
            logger.error(f"Error exporting evidence to Word: {str(e)}")
            raise
