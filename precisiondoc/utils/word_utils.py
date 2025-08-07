import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
from config.config import WORD_FONT_SETTINGS, TABLE_SETTINGS
from .log_utils import setup_logger

# Setup logger for this module
logger = setup_logger(__name__)

class WordUtils:
    """Word document processing utility class"""
    
    @staticmethod
    def apply_word_document_format(doc):
        """
        Apply uniform formatting to Word document
        
        Args:
            doc: Word document object
        """
        # Set page margins and width
        sections = doc.sections
        for section in sections:
            section.page_width = Inches(8.27)  # A4 width
            section.page_height = Inches(11.69)  # A4 height
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            # Default to portrait orientation
            # section.orientation = WD_ORIENT.PORTRAIT  # or WD_ORIENT.LANDSCAPE
            section.orientation = WD_ORIENT.LANDSCAPE

        # Add heading
        heading = doc.add_heading('Precision Evidence Report', 0)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(24)
        heading_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # Set heading font
        for run in heading.runs:
            run.font.name = WORD_FONT_SETTINGS['heading_font_name']
            run.font.size = Pt(WORD_FONT_SETTINGS['heading_font_size'])
            run.font.bold = True
        
        # Add page numbers
        WordUtils.add_page_numbers(doc)
    
    @staticmethod
    def add_page_numbers(doc):
        """
        Add page numbers in 'current/total' format to the document footer
        
        Args:
            doc: Word document object
        """
        # Add page numbers to each section
        for section in doc.sections:
            footer = section.footer
            
            # Clear any existing content in the footer
            for paragraph in footer.paragraphs:
                p = paragraph._p
                p.getparent().remove(p)
            
            # Add a paragraph to the footer
            paragraph = footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add page number field codes
            run = paragraph.add_run()
            
            # Add the 'current page' field
            fld_char1 = OxmlElement('w:fldChar')
            fld_char1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld_char1)
            
            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = ' PAGE '
            run._r.append(instr_text)
            
            fld_char2 = OxmlElement('w:fldChar')
            fld_char2.set(qn('w:fldCharType'), 'end')
            run._r.append(fld_char2)
            
            # Add the separator
            run.add_text(' / ')
            
            # Add the 'total pages' field
            fld_char3 = OxmlElement('w:fldChar')
            fld_char3.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld_char3)
            
            instr_text2 = OxmlElement('w:instrText')
            instr_text2.set(qn('xml:space'), 'preserve')
            instr_text2.text = ' NUMPAGES '
            run._r.append(instr_text2)
            
            fld_char4 = OxmlElement('w:fldChar')
            fld_char4.set(qn('w:fldCharType'), 'end')
            run._r.append(fld_char4)
            
            # Apply formatting to the page number text
            run.font.size = Pt(10)
            run.font.name = WORD_FONT_SETTINGS['page_number_font_name']
    
    @staticmethod
    def apply_paragraph_format(paragraph):
        """
        Apply paragraph formatting
        
        Args:
            paragraph: Paragraph object
        """
        p_format = paragraph.paragraph_format
        # set line spacing to single
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Single line spacing
        # # set line spacing to 1.5
        # p_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5 line spacing
        # # set line spacing to 1.3
        # p_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        # p_format.line_spacing = 1.3  # 1.3 line spacing
        p_format.space_before = Pt(12)  # Space before paragraph
        p_format.space_after = Pt(12)   # Space after paragraph
    
    @staticmethod
    def apply_run_format(run):
        """
        Apply formatting to text run object
        
        Args:
            run: Text run object
        """
        run.font.name = WORD_FONT_SETTINGS['run_font_name']  # Set font to Microsoft YaHei
        run.font.size = Pt(WORD_FONT_SETTINGS['run_font_size'])     # Set font size
    
    @staticmethod
    def apply_separator_format(paragraph):
        """
        Apply formatting to separator paragraph
        
        Args:
            paragraph: Paragraph object
        """
        separator_format = paragraph.paragraph_format
        separator_format.space_before = Pt(12)
        separator_format.space_after = Pt(12)
        
        # Set separator line font
        for run in paragraph.runs:
            run.font.name = WORD_FONT_SETTINGS['separator_font_name']
    
    @staticmethod
    def set_section_orientation(section, orientation='portrait'):
        """
        Set the orientation for a specific section
        
        Args:
            section: Section object
            orientation: 'portrait' or 'landscape'
        """
        if orientation.lower() == 'portrait':
            section.orientation = WD_ORIENT.PORTRAIT
            # When changing orientation, we need to swap width and height
            section.page_width = Inches(8.27)  # A4 width
            section.page_height = Inches(11.69)  # A4 height
        elif orientation.lower() == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            # When changing orientation, we need to swap width and height
            section.page_width = Inches(11.69)  # A4 height becomes width
            section.page_height = Inches(8.27)  # A4 width becomes height
    
    @staticmethod
    def add_text_to_cell(cell, text, multi_line=True):
        """
        Add text to a table cell, either as a single paragraph or multiple paragraphs
        
        Args:
            cell: Table cell object
            text: Text to add
            multi_line: If True, split text by newlines and create a paragraph for each line
        """
        # Clear default paragraph if it exists and is empty
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if not p.runs:  # Only clear if it's empty
                p._element.getparent().remove(p._element)
                p._p = None
                p._element = None
        
        if multi_line:
            # Split text by lines and create a paragraph for each line
            for line in text.strip().split('\n'):
                if line:
                    p = cell.add_paragraph()
                    p.alignment = 0  # Left alignment
                    run = p.add_run(line)
                    WordUtils.apply_run_format(run)
        else:
            # Add all text as a single paragraph
            p = cell.add_paragraph()
            p.alignment = 0  # Left alignment
            run = p.add_run(text)
            WordUtils.apply_run_format(run)
    
    @staticmethod
    def _prepare_evidence_text(row, exclude_columns):
        """
        Prepare evidence text from row data
        
        Args:
            row: DataFrame row
            exclude_columns: Columns to exclude from text
            
        Returns:
            dict: Dictionary containing evidence fields
        """
        evidence_dict = {}
        for col in row.index:
            if col not in exclude_columns:
                # Check if value is empty, if so display "N/A"
                if pd.isna(value := row[col]):
                    evidence_dict[col] = "N/A"
                else:
                    evidence_dict[col] = json.dumps(value, indent=4, ensure_ascii=False)
        return evidence_dict
    
    @staticmethod
    def _create_evidence_table(doc, evidence_dict, multi_line_text=True, show_borders=True):
        """
        Create a table for evidence with text and image placeholders
        
        Args:
            doc: Word document
            evidence_dict: Dictionary containing evidence fields
            multi_line_text: Whether to split text into multiple lines (True = multi-row table, False = JSON dict)
            show_borders: Whether to show table borders
            
        Returns:
            tuple: (table, left_cells, right_cell)
        """
        if multi_line_text:
            # Create a table with multiple rows (one row per dictionary item) and 2 columns
            row_count = len(evidence_dict)
            table = doc.add_table(rows=row_count, cols=2)
        else:
            # Create a table with 1 row and 2 columns for JSON format
            table = doc.add_table(rows=1, cols=2)
            
        table.autofit = False
        table.allow_autofit = False
        
        # Determine page orientation and set table width accordingly
        # Get the current section
        current_section = doc.sections[-1]
        is_landscape = current_section.orientation == WD_ORIENT.LANDSCAPE
        
        # Set table width based on orientation (accounting for margins) using config settings
        if is_landscape:
            table_width = Inches(TABLE_SETTINGS['landscape_table_width'])
        else:
            table_width = Inches(TABLE_SETTINGS['portrait_table_width'])
            
        # Calculate column widths based on ratio settings from config
        text_col_width = Inches(table_width.inches * TABLE_SETTINGS['text_column_ratio'])
        image_col_width = Inches(table_width.inches * TABLE_SETTINGS['image_column_ratio'])
            
        table.width = table_width
        
        # Set column widths using direct XML manipulation for more reliable width control
        tbl_grid = table._element.find(".//w:tblGrid", table._element.nsmap)
        
        # Clear existing grid columns if any
        for grid_col in tbl_grid.findall(".//w:gridCol", table._element.nsmap):
            tbl_grid.remove(grid_col)
        
        # Create new grid columns with desired widths
        col1 = OxmlElement('w:gridCol')
        col1.set(qn('w:w'), str(int(text_col_width.twips)))  # Text column - 20%
        tbl_grid.append(col1)
        
        col2 = OxmlElement('w:gridCol')
        col2.set(qn('w:w'), str(int(image_col_width.twips)))  # Image column - 80%
        tbl_grid.append(col2)
        
        # Set column widths at the cell level too
        for row in table.rows:
            row.cells[0].width = text_col_width
            row.cells[1].width = image_col_width
        
        # Set table borders
        if show_borders:
            # Show all borders
            table.style = 'Table Grid'
        else:
            # Hide all borders
            table.style = 'Table Normal'
            # Ensure no borders are visible
            for row in table.rows:
                for cell in row.cells:
                    # Set border width to 0 using XML directly
                    tcPr = cell._element.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        cell._element.append(tcPr)
                    
                    # Create border element for each side
                    for side in ['top', 'left', 'bottom', 'right']:
                        # Create border element if it doesn't exist
                        tc_borders = tcPr.find('.//w:tcBorders', namespaces=tcPr.nsmap)
                        if tc_borders is None:
                            tc_borders = OxmlElement('w:tcBorders')
                            tcPr.append(tc_borders)
                        
                        # Create or find the specific border
                        border_elem_tag = f'w:{side}'
                        border = tc_borders.find(border_elem_tag, namespaces=tcPr.nsmap)
                        if border is None:
                            border = OxmlElement(border_elem_tag)
                            tc_borders.append(border)
                        
                        # Set border to none
                        border.set(qn('w:val'), 'nil')
        
        # Set table properties for continuation across pages
        # This allows the table to break across pages if needed
        table_pr = table._element.xpath('w:tblPr')[0]
        table_layout = OxmlElement('w:tblLayout')
        table_layout.set(qn('w:type'), 'fixed')
        table_pr.append(table_layout)
        
        # Allow row to break across pages
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), '0')  # 0 means can split
            trPr.append(cantSplit)
        
        left_cells = []
        
        if multi_line_text:
            # Add text content to the left cells - one row per dictionary item
            for i, (key, value) in enumerate(evidence_dict.items()):
                if i < len(table.rows):  # Safety check
                    left_cell = table.cell(i, 0)
                    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    p = left_cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(f"{key}: {value}")
                    run.font.name = WORD_FONT_SETTINGS['run_font_name']  # Use Microsoft YaHei font for both English and Chinese text
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), WORD_FONT_SETTINGS['run_font_name'])  # 设置东亚字体
                    run.font.size = Pt(WORD_FONT_SETTINGS['run_font_size'])
                    left_cells.append(left_cell)
                
            # Merge cells in the right column
            right_cell = table.cell(0, 1)
            if row_count > 1:
                # Merge cells vertically in the right column
                for i in range(1, row_count):
                    right_cell.merge(table.cell(i, 1))
        else:
            # Format the evidence dictionary into a single JSON-style text
            left_cell = table.cell(0, 0)
            left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Create a properly formatted dictionary string without double quotes
            formatted_dict = "{\n"
            for i, (key, value) in enumerate(evidence_dict.items()):
                # Format the value based on its type
                if isinstance(value, str):
                    formatted_value = f'{value}'
                else:
                    formatted_value = str(value)
                    
                # Add comma for all items except the last one
                if i < len(evidence_dict) - 1:
                    formatted_dict += f'    "{key}": {formatted_value},\n'
                else:
                    formatted_dict += f'    "{key}": {formatted_value}\n'
            
            formatted_dict += "}"
            
            # Add the formatted dictionary string as a single paragraph with preserved formatting
            p = left_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(formatted_dict)
            run.font.name = WORD_FONT_SETTINGS['run_font_name']  # Use Microsoft YaHei font for both English and Chinese text
            run._element.rPr.rFonts.set(qn('w:eastAsia'), WORD_FONT_SETTINGS['run_font_name'])  # Set East Asia font
            run.font.size = Pt(WORD_FONT_SETTINGS['run_font_size'])
            
            left_cells = [left_cell]
            right_cell = table.cell(0, 1)
        
        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        return table, left_cells, right_cell
    
    @staticmethod
    def _add_image_to_cell(cell, row, output_folder, doc=None):
        """
        Add image to table cell
        
        Args:
            cell: Table cell to add image to
            row: DataFrame row with image information
            output_folder: Folder containing images
            doc: Word document (for determining orientation)
            
        Returns:
            bool: True if image was found and added, False otherwise
        """
        img_found = False
        
        # Calculate image width based on cell width (slightly smaller to allow for margins)
        # Get the image width from the TABLE_SETTINGS or use 90% of the image column width
        image_width_setting = TABLE_SETTINGS.get('image_width')
        if image_width_setting is not None:
            image_width = Inches(image_width_setting)
        else:
            # Determine which table width to use based on orientation
            if doc and doc.sections and len(doc.sections) > 0:
                current_section = doc.sections[-1]
                is_landscape = current_section.orientation == WD_ORIENT.LANDSCAPE
                if is_landscape:
                    table_width = TABLE_SETTINGS['landscape_table_width']
                else:
                    table_width = TABLE_SETTINGS['portrait_table_width']
            else:
                # Default to portrait if we can't determine orientation
                table_width = TABLE_SETTINGS['portrait_table_width']
                
            # Calculate based on table width and image column ratio
            calculated_width = table_width * TABLE_SETTINGS['image_column_ratio'] * 0.9
            image_width = Inches(calculated_width)
        
        # First check if image_path is directly provided
        if 'image_path' in row and not pd.isna(row['image_path']):
            img_path = row['image_path']
            if os.path.exists(img_path):
                try:
                    img_p = cell.paragraphs[0]
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
                    img_run = img_p.add_run()
                    img_run.add_picture(img_path, width=image_width)  # Dynamic width based on cell
                    img_found = True
                except Exception as e:
                    logger.warning(f"Failed to add image {img_path}: {str(e)}")
        
        # If image_path not provided or invalid, try to find image based on page number
        if not img_found and 'page_number' in row and not pd.isna(row['page_number']):
            page_num = int(row['page_number'])
            doc_type = row.get('document_type', '')
            
            # Build possible image paths
            img_folder = os.path.join(output_folder, 'images')
            possible_paths = [
                os.path.join(img_folder, f"{doc_type}_page_{page_num}.png"),
                os.path.join(img_folder, f"page_{page_num}.png"),
                os.path.join(output_folder, f"{doc_type}_page_{page_num}.png"),
                os.path.join(output_folder, f"page_{page_num}.png")
            ]
            
            for img_path in possible_paths:
                if os.path.exists(img_path):
                    try:
                        # Add image to the right cell
                        img_p = cell.paragraphs[0]
                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
                        img_run = img_p.add_run()
                        img_run.add_picture(img_path, width=image_width)  # Dynamic width based on cell
                        img_found = True
                        break
                    except Exception as e:
                        logger.warning(f"Failed to add image {img_path}: {str(e)}")
        
        if not img_found:
            logger.warning(f"No image found for page {int(row['page_number']) if 'page_number' in row and not pd.isna(row['page_number']) else 'unknown'} of document {row.get('document_type', '')}")
        
        return img_found
    
    @staticmethod
    def _add_separator_and_section(doc, idx):
        """
        Add separator line and create new section if needed
        
        Args:
            doc: Word document
            idx: Current row index
        """
        # Add separator line
        separator = doc.add_paragraph('-' * 50)
        WordUtils.apply_separator_format(separator)
        
        # Create a new section for next evidence item
        if idx > 0:
            return doc.add_section()
        return None
    
    @staticmethod
    def export_evidence_to_word(excel_file, word_file, output_folder, multi_line_text=True, show_borders=True):
        """
        Export precision evidence from Excel to Word document
        
        Args:
            excel_file: Path to Excel file
            word_file: Path to output Word file
            output_folder: Output folder path, used to find images
            multi_line_text: If True, split text by newlines in the left cell
            show_borders: If True, show table borders
        """
        try:
            logger.info(f"Exporting evidence from {excel_file} to {word_file}")
            
            # Read Excel file
            df = pd.read_excel(excel_file)
            
            # Filter precision evidence
            if 'is_precision_evidence' in df.columns:
                evidence_df = df[df['is_precision_evidence'] == True].copy()
                
                if evidence_df.empty:
                    logger.warning("No precision evidence found in Excel file")
                    return
                
                # Create Word document
                doc = Document()
                
                # Apply document formatting
                WordUtils.apply_word_document_format(doc)
                
                # Define columns to exclude from evidence text
                exclude_columns = (
                    'page_type', 'is_precision_evidence',
                    'page_number', 'success', 'document_type', '解析', '分析', '结论', '文字提取', 'evidence_level'
                )
                
                # Process evidence row by row
                for idx, row in evidence_df.iterrows():
                    # Create a new section for each evidence item to control page layout
                    if idx > 0:
                        new_section = doc.add_section()
                        # By default, use landscape orientation
                        WordUtils.set_section_orientation(new_section, 'landscape')
                    
                    # Prepare evidence text
                    evidence_dict = WordUtils._prepare_evidence_text(row, exclude_columns)
                    
                    # Create table with text and image placeholders
                    _, left_cells, right_cell = WordUtils._create_evidence_table(doc, evidence_dict, 
                                                                       multi_line_text=multi_line_text,
                                                                       show_borders=show_borders)
                    
                    # Add image to right cell
                    WordUtils._add_image_to_cell(right_cell, row, output_folder, doc)
                    
                    # Add separator line
                    separator = doc.add_paragraph('-' * 50)
                    WordUtils.apply_separator_format(separator)
                
                # Save Word document
                doc.save(word_file)
                logger.info(f"Evidence exported to Word file: {word_file}")
            else:
                logger.warning("No 'is_precision_evidence' column found in Excel file")
        except Exception as e:
            logger.error(f"Error exporting evidence to Word: {str(e)}")
            raise
